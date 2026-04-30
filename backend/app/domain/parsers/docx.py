# Requires: python-docx
# Install with: pip install python-docx

from typing import TYPE_CHECKING

from app.domain.parsers.base import ParsedBlock, ParsedDocument
from app.domain.parsers.failures import ParseFailure

if TYPE_CHECKING:
    import docx


class DocxParser:
    supported_extensions = (".docx",)
    parser_name = "docx"
    parser_version = "docx_parser_v1"

    def __init__(self) -> None:
        """Initialize the DOCX parser, checking for required dependencies."""
        try:
            import docx  # noqa: F401
        except ImportError as e:
            self._import_error = e
        else:
            self._import_error = None

    def parse(self, content: bytes) -> ParsedDocument | ParseFailure:
        """
        Extract paragraphs and tables from DOCX file.
        Groups into sections by Heading styles.
        Converts tables to text (tab-separated cells, newline-separated rows).
        """
        if self._import_error:
            raise ImportError(
                "python-docx is required to parse DOCX files. "
                "Install with: pip install python-docx"
            ) from self._import_error

        import io

        try:
            from docx import Document
        except ImportError as e:
            raise ImportError(
                "python-docx is required. Install with: pip install python-docx"
            ) from e

        try:
            doc = Document(io.BytesIO(content))
        except Exception as e:
            return ParseFailure(
                filename="",
                reason="corrupt",
                error_message=f"Failed to parse DOCX: {str(e)}",
            )

        if not doc.paragraphs and not doc.tables:
            return ParseFailure(
                filename="",
                reason="empty",
                error_message="DOCX file has no content",
            )

        blocks: list[ParsedBlock] = []
        section_stack: list[tuple[int, str]] = []
        normalized_lines: list[str] = []

        def current_section_path() -> tuple[str, ...]:
            return tuple(heading for _, heading in section_stack)

        for element in doc.element.body:
            # Process paragraphs
            if element.tag.endswith("p"):
                # Map to paragraph object
                para = next(
                    (p for p in doc.paragraphs if p._element is element), None
                )
                if not para:
                    continue

                text = para.text.strip()
                if not text:
                    continue

                style_name = para.style.name if para.style else ""

                # Check if this is a heading
                if style_name.startswith("Heading"):
                    # Extract level from style name (e.g., "Heading 1" -> level 1)
                    level = 1
                    if "Heading" in style_name:
                        try:
                            level = int(style_name.split()[-1])
                        except (ValueError, IndexError):
                            level = 1

                    # Update section stack
                    section_stack[:] = [
                        (existing_level, value)
                        for existing_level, value in section_stack
                        if existing_level < level
                    ]
                    section_stack.append((level, text))

                    blocks.append(
                        ParsedBlock(
                            kind="heading",
                            text=text,
                            section_path=current_section_path(),
                            heading=text,
                            level=level,
                        )
                    )
                    normalized_lines.append(text)
                else:
                    # Regular paragraph
                    blocks.append(
                        ParsedBlock(
                            kind="paragraph",
                            text=text,
                            section_path=current_section_path(),
                        )
                    )
                    normalized_lines.append(text)

            # Process tables
            elif element.tag.endswith("tbl"):
                # Map to table object
                table = next(
                    (t for t in doc.tables if t._element is element), None
                )
                if not table:
                    continue

                table_text = _convert_table_to_text(table)
                if table_text:
                    blocks.append(
                        ParsedBlock(
                            kind="table",
                            text=table_text,
                            section_path=current_section_path(),
                        )
                    )
                    normalized_lines.append(table_text)

        normalized_text = "\n".join(normalized_lines)
        if normalized_text:
            normalized_text += "\n"

        return ParsedDocument(
            normalized_text=normalized_text,
            blocks=tuple(blocks),
            parser_name=self.parser_name,
            parser_version=self.parser_version,
        )


def _convert_table_to_text(table) -> str:
    """Convert a DOCX table to tab-separated text."""
    rows: list[str] = []

    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        rows.append("\t".join(cells))

    return "\n".join(rows)
