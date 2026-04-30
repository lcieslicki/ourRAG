import io

import pytest

from app.domain.parsers.docx import DocxParser
from app.domain.parsers.failures import ParseFailure


def test_parse_empty_docx_returns_minimal_result() -> None:
    """Test parsing an empty DOCX file."""
    try:
        from docx import Document
    except ImportError:
        pytest.skip("python-docx not installed")

    parser = DocxParser()

    # Create empty DOCX in memory
    doc = Document()
    docx_bytes = io.BytesIO()
    doc.save(docx_bytes)
    content = docx_bytes.getvalue()

    result = parser.parse(content)

    # Empty DOCX should return ParseFailure with reason "empty"
    assert isinstance(result, ParseFailure)
    assert result.reason == "empty"


def test_parse_docx_with_headings_creates_sections() -> None:
    """Test that Heading styles create section blocks."""
    try:
        from docx import Document
    except ImportError:
        pytest.skip("python-docx not installed")

    parser = DocxParser()

    # Create DOCX with headings and paragraphs
    doc = Document()
    heading1 = doc.add_heading("Main Section", level=1)
    para1 = doc.add_paragraph("Paragraph under main section")
    heading2 = doc.add_heading("Subsection", level=2)
    para2 = doc.add_paragraph("Paragraph under subsection")

    docx_bytes = io.BytesIO()
    doc.save(docx_bytes)
    content = docx_bytes.getvalue()

    result = parser.parse(content)

    assert not isinstance(result, ParseFailure)
    assert result.parser_name == "docx"

    # Should have heading and paragraph blocks
    blocks = result.blocks
    kinds = [b.kind for b in blocks]
    assert "heading" in kinds
    assert "paragraph" in kinds

    # Check headings
    heading_blocks = [b for b in blocks if b.kind == "heading"]
    assert len(heading_blocks) >= 1
    assert heading_blocks[0].heading == "Main Section"


def test_parse_corrupt_bytes_returns_failure() -> None:
    parser = DocxParser()

    content = b"not a valid docx"
    result = parser.parse(content)

    assert isinstance(result, ParseFailure)
    assert result.reason == "corrupt"


def test_docx_parser_handles_import_error_gracefully() -> None:
    """Test that missing python-docx dependency raises informative error."""
    parser = DocxParser()

    # Simulate import error
    parser._import_error = ImportError("No module named 'docx'")

    content = b"some content"

    with pytest.raises(ImportError) as exc_info:
        parser.parse(content)

    assert "python-docx is required" in str(exc_info.value)
    assert "pip install python-docx" in str(exc_info.value)
